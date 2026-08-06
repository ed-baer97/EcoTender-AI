"""Deterministic document text extraction for LLM evidence packs.

Extract once at ingest (or lazily from MinIO), cache excerpts by sha256.
Never send raw PDF bytes to the LLM.
"""

from __future__ import annotations

import io
import logging
import os
import re
from datetime import datetime, timezone
from typing import Any

logger = logging.getLogger("ecotender.doc_extract")

MAX_EXCERPT_CHARS = int(os.getenv("DOC_EXTRACT_MAX_CHARS", "10000"))
MAX_FILE_BYTES = int(os.getenv("DOC_EXTRACT_MAX_BYTES", str(8 * 1024 * 1024)))
DOC_KINDS = frozenset({"document", "specification"})


def _clean_text(text: str, max_chars: int = MAX_EXCERPT_CHARS) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = text.strip()
    if len(text) > max_chars:
        return text[: max_chars - 1].rstrip() + "…"
    return text


def extract_text_from_bytes(
    data: bytes,
    *,
    content_type: str = "",
    filename: str = "",
    max_chars: int = MAX_EXCERPT_CHARS,
) -> tuple[str, str | None]:
    """Return (excerpt, error). Empty excerpt with error on failure."""
    if not data:
        return "", "empty_bytes"
    if len(data) > MAX_FILE_BYTES:
        data = data[:MAX_FILE_BYTES]

    name = (filename or "").lower()
    ctype = (content_type or "").lower()
    try:
        if "pdf" in ctype or name.endswith(".pdf"):
            return _extract_pdf(data, max_chars), None
        if (
            "wordprocessingml" in ctype
            or "msword" in ctype
            or name.endswith(".docx")
            or name.endswith(".doc")
        ):
            if name.endswith(".doc") and not name.endswith(".docx"):
                return "", "unsupported_doc_legacy"
            return _extract_docx(data, max_chars), None
        if "html" in ctype or name.endswith((".html", ".htm")):
            return _extract_html(data, max_chars), None
        if "text/" in ctype or name.endswith((".txt", ".csv", ".rtf")):
            return _clean_text(data.decode("utf-8", errors="replace"), max_chars), None
        # sniff PDF magic
        if data[:4] == b"%PDF":
            return _extract_pdf(data, max_chars), None
        if data[:2] == b"PK":
            return _extract_docx(data, max_chars), None
        return "", f"unsupported_type:{ctype or name or 'unknown'}"
    except Exception as exc:  # noqa: BLE001
        logger.warning("extract failed name=%s err=%s", filename, exc)
        return "", f"extract_failed:{type(exc).__name__}"


def _extract_pdf(data: bytes, max_chars: int) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    total = 0
    for page in reader.pages:
        chunk = page.extract_text() or ""
        if not chunk.strip():
            continue
        parts.append(chunk)
        total += len(chunk)
        if total >= max_chars:
            break
    return _clean_text("\n".join(parts), max_chars)


def _extract_docx(data: bytes, max_chars: int) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
        if sum(len(p) for p in parts) >= max_chars:
            break
    return _clean_text("\n".join(parts), max_chars)


def _extract_html(data: bytes, max_chars: int) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(data.decode("utf-8", errors="replace"), "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return _clean_text(soup.get_text("\n"), max_chars)


def build_extract_record(
    *,
    name: str,
    kind: str,
    sha256: str | None,
    excerpt: str,
    error: str | None = None,
    content_type: str | None = None,
    object_key: str | None = None,
    group_name: str | None = None,
) -> dict[str, Any]:
    return {
        "name": name,
        "kind": kind,
        "sha256": sha256,
        "excerpt": excerpt,
        "chars": len(excerpt or ""),
        "truncated": bool(excerpt and excerpt.endswith("…")),
        "content_type": content_type,
        "object_key": object_key,
        "group_name": group_name,
        "extracted_at": datetime.now(timezone.utc).isoformat(),
        **({"error": error} if error else {}),
    }


def merge_doc_extracts(
    existing: list[dict[str, Any]] | None,
    new_items: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Keep prior extracts for unchanged sha256; replace/add new ones."""
    by_key: dict[str, dict[str, Any]] = {}
    for item in existing or []:
        key = str(item.get("sha256") or item.get("object_key") or item.get("name") or "")
        if key:
            by_key[key] = item
    for item in new_items:
        key = str(item.get("sha256") or item.get("object_key") or item.get("name") or "")
        if not key:
            continue
        prev = by_key.get(key)
        # Skip re-extract if we already have a non-empty excerpt for same hash.
        if prev and prev.get("sha256") and prev.get("sha256") == item.get("sha256") and prev.get("excerpt") and not item.get("error"):
            continue
        if prev and prev.get("excerpt") and item.get("error") and not item.get("excerpt"):
            continue
        by_key[key] = item
    return list(by_key.values())


def rank_doc_extracts(
    extracts: list[dict[str, Any]] | None,
    *,
    limit: int = 5,
    excerpt_chars: int = 3500,
) -> list[dict[str, Any]]:
    """Prefer specification / ТЗ docs, then ones with text."""

    def sort_key(item: dict[str, Any]) -> tuple[int, int, str]:
        kind = str(item.get("kind") or "")
        group = str(item.get("group_name") or "").lower()
        name = str(item.get("name") or "").lower()
        is_spec = (
            kind == "specification"
            or "спецификац" in group
            or "спецификац" in name
            or "техн" in group
            or "тз" in name
        )
        has_text = 1 if item.get("excerpt") else 0
        return (0 if is_spec else 1, 0 if has_text else 1, name)

    ranked = sorted(extracts or [], key=sort_key)
    out: list[dict[str, Any]] = []
    for item in ranked[:limit]:
        excerpt = str(item.get("excerpt") or "")
        if len(excerpt) > excerpt_chars:
            excerpt = excerpt[: excerpt_chars - 1].rstrip() + "…"
        out.append(
            {
                "name": item.get("name"),
                "kind": item.get("kind"),
                "group_name": item.get("group_name"),
                "chars": item.get("chars"),
                "truncated": item.get("truncated") or (excerpt.endswith("…") if excerpt else False),
                "excerpt": excerpt,
                **({"error": item["error"]} if item.get("error") else {}),
            }
        )
    return out


def fetch_minio_bytes(bucket: str, object_key: str) -> bytes:
    from minio import Minio

    endpoint = os.getenv("MINIO_ENDPOINT", "minio:9000")
    access_key = os.getenv("MINIO_ACCESS_KEY", "minioadmin")
    secret_key = os.getenv("MINIO_SECRET_KEY", "minioadmin")
    secure = os.getenv("MINIO_SECURE", "false").lower() in {"1", "true", "yes", "on"}
    client = Minio(endpoint, access_key=access_key, secret_key=secret_key, secure=secure)
    response = client.get_object(bucket, object_key)
    try:
        return response.read(MAX_FILE_BYTES + 1)
    finally:
        response.close()
        response.release_conn()


def ensure_doc_extracts_from_minio(gos: dict[str, Any]) -> list[dict[str, Any]]:
    """Lazy-fill missing extracts from MinIO for documents with object_key."""
    existing = list(gos.get("doc_extracts") or [])
    have = {str(x.get("sha256") or x.get("object_key")) for x in existing if x.get("excerpt")}
    new_items: list[dict[str, Any]] = []
    for doc in gos.get("documents") or []:
        kind = str(doc.get("kind") or "document")
        if kind not in DOC_KINDS and "document" not in kind:
            continue
        sha = doc.get("sha256")
        key = doc.get("object_key")
        bucket = doc.get("bucket") or os.getenv("MINIO_RAW_BUCKET", "ecotender-raw")
        lookup = str(sha or key or "")
        if not key or lookup in have:
            continue
        try:
            blob = fetch_minio_bytes(str(bucket), str(key))
            excerpt, err = extract_text_from_bytes(
                blob,
                content_type=str(doc.get("content_type") or ""),
                filename=str(doc.get("name") or key),
            )
            new_items.append(
                build_extract_record(
                    name=str(doc.get("name") or key),
                    kind=kind,
                    sha256=sha,
                    excerpt=excerpt,
                    error=err,
                    content_type=doc.get("content_type"),
                    object_key=key,
                    group_name=doc.get("group_name"),
                )
            )
        except Exception as exc:  # noqa: BLE001
            new_items.append(
                build_extract_record(
                    name=str(doc.get("name") or key),
                    kind=kind,
                    sha256=sha,
                    excerpt="",
                    error=f"minio_fetch_failed:{exc}",
                    content_type=doc.get("content_type"),
                    object_key=key,
                    group_name=doc.get("group_name"),
                )
            )
    return merge_doc_extracts(existing, new_items)
